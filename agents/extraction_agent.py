"""
Agent 1: Extraction Agent
Handles PDF (text + OCR fallback) and DOCX extraction.
"""
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class ExtractionAgent:
    """Extracts raw text from PDF or DOCX contracts."""

    OCR_THRESHOLD_CHARS_PER_PAGE = 100

    def extract(self, file_path: str) -> dict:
        path = Path(file_path)
        ext = path.suffix.lower()
        file_size = os.path.getsize(file_path)

        if ext == ".pdf":
            return self._extract_pdf(file_path, path.name, file_size)
        elif ext in (".docx", ".doc"):
            return self._extract_docx(file_path, path.name, file_size)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Expected .pdf or .docx")

    # ------------------------------------------------------------------
    # PDF extraction
    # ------------------------------------------------------------------

    def _extract_pdf(self, file_path: str, file_name: str, file_size: int) -> dict:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError("PyMuPDF (fitz) not installed. Run: pip install pymupdf")

        doc = fitz.open(file_path)
        pages = len(doc)
        page_texts = []

        for page_num in range(pages):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            page_texts.append(text)

        total_chars = sum(len(t) for t in page_texts)
        avg_chars_per_page = total_chars / max(pages, 1)

        if avg_chars_per_page < self.OCR_THRESHOLD_CHARS_PER_PAGE:
            logger.info(
                f"PDF appears scanned (avg {avg_chars_per_page:.1f} chars/page). "
                "Falling back to OCR."
            )
            doc.close()
            return self._extract_pdf_ocr(file_path, file_name, file_size, pages)

        raw_text = "\n\n".join(
            f"[PAGE {i + 1}]\n{text}" for i, text in enumerate(page_texts)
        )
        doc.close()

        return {
            "raw_text": raw_text.strip(),
            "pages": pages,
            "extraction_method": "pymupdf",
            "file_name": file_name,
            "file_size_bytes": file_size,
        }

    def _extract_pdf_ocr(
        self, file_path: str, file_name: str, file_size: int, pages: int
    ) -> dict:
        try:
            import fitz
            from PIL import Image
            import pytesseract
            import io
        except ImportError as e:
            raise RuntimeError(
                f"OCR dependencies not installed: {e}. "
                "Run: pip install pytesseract Pillow"
            )

        doc = fitz.open(file_path)
        page_texts = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            # Render at 200 DPI for decent OCR quality
            mat = fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))
            text = pytesseract.image_to_string(img, config="--psm 6")
            page_texts.append(text)
            logger.debug(f"OCR page {page_num + 1}: {len(text)} chars")

        doc.close()
        raw_text = "\n\n".join(
            f"[PAGE {i + 1}]\n{text}" for i, text in enumerate(page_texts)
        )

        return {
            "raw_text": raw_text.strip(),
            "pages": pages,
            "extraction_method": "ocr",
            "file_name": file_name,
            "file_size_bytes": file_size,
        }

    # ------------------------------------------------------------------
    # DOCX extraction
    # ------------------------------------------------------------------

    def _extract_docx(self, file_path: str, file_name: str, file_size: int) -> dict:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise RuntimeError(
                "python-docx not installed. Run: pip install python-docx"
            )

        doc = Document(file_path)
        sections = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            # Preserve heading hierarchy
            if "Heading" in style_name:
                level = self._extract_heading_level(style_name)
                prefix = "#" * level + " "
                sections.append(prefix + para.text.strip())
            else:
                sections.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        raw_text = "\n\n".join(sections)

        return {
            "raw_text": raw_text.strip(),
            "pages": self._estimate_pages_from_text(raw_text),
            "extraction_method": "docx",
            "file_name": file_name,
            "file_size_bytes": file_size,
        }

    @staticmethod
    def _extract_heading_level(style_name: str) -> int:
        """Extract numeric level from style like 'Heading 1', 'Heading 2'."""
        parts = style_name.split()
        if len(parts) >= 2 and parts[-1].isdigit():
            return int(parts[-1])
        return 1

    @staticmethod
    def _estimate_pages_from_text(text: str) -> int:
        """Rough estimate: ~3000 chars per page."""
        return max(1, len(text) // 3000)
